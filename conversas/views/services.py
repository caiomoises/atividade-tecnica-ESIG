import os
import tempfile
from pypdf import PdfReader
from docx import Document as DocxDocument
from ..models.vetorRAG import VetorDocumento
from ..models import Conversa, Mensagem, Documento
from django.utils import timezone
from langchain_community.embeddings import OpenAIEmbeddings
from langchain.schema import Document
from decouple import config
import logging
import openai
from pgvector.django import L2Distance

logger = logging.getLogger(__name__)

def extrair_texto(arquivo):
    """Extrai texto de um arquivo PDF ou DOCX com tratamento de erros"""
    ext = os.path.splitext(arquivo.name)[-1].lower()
    texto = ""

    try:
        with tempfile.NamedTemporaryFile(delete=False) as temp_file:
            for chunk in arquivo.chunks():
                temp_file.write(chunk)
            temp_file_path = temp_file.name

        if ext == ".pdf":
            reader = PdfReader(temp_file_path)
            for page in reader.pages:
                texto += page.extract_text() or "" + "\n"

        elif ext == ".docx":
            doc = DocxDocument(temp_file_path)
            for para in doc.paragraphs:
                texto += para.text + "\n"

        return texto.strip()

    except Exception as e:
        logger.error(f"Erro ao extrair texto do arquivo: {e}")
        raise ValueError(f"Erro ao processar arquivo: {str(e)}")
    
    finally:
        if temp_file_path and os.path.exists(temp_file_path):
            os.remove(temp_file_path)

def limpar_embeddings_documento(documento_id):
    """Remove todos os embeddings associados a um documento"""
    try:
        deleted_count, _ = VetorDocumento.objects.filter(documento_id=documento_id).delete()
        logger.info(f"Removidos {deleted_count} embeddings para documento {documento_id}")
        return deleted_count
    except Exception as e:
        logger.error(f"Erro ao limpar embeddings: {e}")
        raise

def gerar_e_salvar_embedding(texto, documento):
    """Gera embeddings do texto em chunks e armazena no PostgreSQL"""
    try:
        # Verifica se o texto está vazio
        if not texto.strip():
            raise ValueError("Texto vazio não pode ser processado")

        # Limpa embeddings antigos primeiro
        limpar_embeddings_documento(documento.id)

        openai_api_key = config("OPENAI_API_KEY")
        if not openai_api_key:
            raise ValueError("OPENAI_API_KEY não configurada")

        embeddings = OpenAIEmbeddings(openai_api_key=openai_api_key)
        
        # Divide o texto em chunks para evitar limite de tokens
        chunk_size = 1000
        chunks = [texto[i:i+chunk_size] for i in range(0, len(texto), chunk_size)]
        
        for i, chunk in enumerate(chunks):
            try:
                embedding = embeddings.embed_documents([chunk])[0]
                
                doc = Document(
                    page_content=chunk,
                    metadata={
                        "documento_id": documento.id,
                        "documento_nome": documento.nome,
                        "chunk_index": i
                    }
                )

                vetor = VetorDocumento(
                    page_content=chunk,
                    documento=documento,
                    embedding=embedding,
                    metadata=doc.metadata
                )
                vetor.save()
                logger.debug(f"Embedding salvo para chunk {i} do documento {documento.id}")

            except Exception as chunk_error:
                logger.error(f"Erro no chunk {i}: {chunk_error}")
                continue

        return len(chunks)

    except Exception as e:
        logger.error(f"Erro ao gerar e salvar embedding: {e}")
        raise

def obter_historico(conversa, limit=10):
    """Retorna o histórico de mensagens no formato do OpenAI"""
    mensagens = conversa.mensagens.all().order_by('-criado_em')[:limit]
    historico = []
    
    for msg in reversed(mensagens):  # Ordena do mais antigo para o mais novo
        historico.append({"role": "user", "content": msg.texto})
        if msg.resposta:
            historico.append({"role": "assistant", "content": msg.resposta})
    
    return historico

def buscar_documentos_relevantes(pergunta, documento_id=None, limit=3):
    """Busca os documentos mais relevantes para a pergunta"""
    try:
        openai_api_key = config("OPENAI_API_KEY")
        embeddings = OpenAIEmbeddings(openai_api_key=openai_api_key)
        
        pergunta_embedding = embeddings.embed_documents([pergunta])[0]
        
        query = VetorDocumento.objects.all()
        
        if documento_id:
            query = query.filter(documento_id=documento_id)
        
        documentos_similares = query.order_by(
            L2Distance('embedding', pergunta_embedding)
        )[:limit]
        
        return list(documentos_similares)
    
    except Exception as e:
        logger.error(f"Erro ao buscar documentos relevantes: {e}")
        return []

def construir_contexto(documentos_similares):
    """Constrói o contexto a partir dos documentos similares"""
    contexto = ""
    for i, doc in enumerate(documentos_similares, 1):
        contexto += (
            f"\n\n[Documento {i}: {doc.documento.nome}]\n"
            f"{doc.page_content[:2000]}"  # Limita o tamanho do chunk
        )
    return contexto.strip()

def enviar_mensagem(pergunta, contexto="", historico=None, model="gpt-4-turbo"):
    """Envia mensagem ao OpenAI com contexto e histórico"""
    try:
        messages = []
        
        # Adiciona contexto como mensagem de sistema
        if contexto:
            messages.append({
                "role": "system",
                "content": (
                    "Você é um assistente especializado em analisar documentos. "
                    "Baseie suas respostas estritamente nestes documentos:\n"
                    f"{contexto}\n\n"
                    "Se a pergunta não estiver relacionada aos documentos, "
                    "responda que não possui informações sobre o assunto."
                )
            })
        
        # Adiciona histórico da conversa
        if historico:
            messages.extend(historico)
        
        # Adiciona a pergunta atual
        messages.append({"role": "user", "content": pergunta})

        resposta = openai.chat.completions.create(
            model=model,
            messages=messages,
            temperature=0.5,  # Menos criatividade para respostas mais precisas
            max_tokens=1500
        )

        return resposta.choices[0].message.content

    except openai.RateLimitError:
        logger.error("Limite de taxa da OpenAI excedido")
        return "Estou recebendo muitas solicitações. Por favor, tente novamente mais tarde."
    
    except Exception as e:
        logger.error(f"Erro ao enviar mensagem: {e}")
        return "Desculpe, ocorreu um erro ao processar sua mensagem."

def processar_arquivo(arquivo, usuario=None):
    """Processa um arquivo enviado pelo usuário"""
    try:
        # Extrai texto do arquivo
        texto = extrair_texto(arquivo)
        
        # Cria registro do documento
        documento = Documento.objects.create(
            nome=arquivo.name,
            usuario=usuario,
            criado_em=timezone.now()
        )
        
        # Gera e salva embeddings
        chunks_processados = gerar_e_salvar_embedding(texto, documento)
        
        return {
            "success": True,
            "documento_id": documento.id,
            "documento_nome": documento.nome,
            "chunks_processados": chunks_processados
        }
    
    except Exception as e:
        logger.error(f"Erro ao processar arquivo: {e}")
        return {
            "success": False,
            "error": str(e)
        }

def buscar_contexto(pergunta, usuario=None, conversa_id=None, documento_id=None):
    """Busca contexto e gera resposta com RAG"""
    try:
        # Configurações iniciais
        openai_api_key = config("OPENAI_API_KEY")
        if not openai_api_key:
            raise ValueError("OPENAI_API_KEY não configurada")
        
        openai.api_key = openai_api_key

        # Gerenciamento da conversa
        if conversa_id:
            conversa = Conversa.objects.get(id=conversa_id)
        else:
            conversa = Conversa.objects.create(
                usuario=usuario.username if usuario else "Anônimo",
                criado_em=timezone.now()
            )

        # Registra a pergunta do usuário
        mensagem = Mensagem.objects.create(
            conversa=conversa,
            texto=pergunta,
            documento_id=documento_id
        )

        try:
            # Busca documentos relevantes
            documentos_similares = buscar_documentos_relevantes(pergunta, documento_id)
            
            # Constrói contexto
            contexto = construir_contexto(documentos_similares)
            
            # Obtém histórico da conversa
            historico = obter_historico(conversa)
            
            # Gera resposta
            resposta = enviar_mensagem(pergunta, contexto, historico)
            
            # Atualiza a mensagem com a resposta
            mensagem.resposta = resposta
            mensagem.save()

            return {
                "success": True,
                "resposta": resposta,
                "conversa_id": conversa.id,
                "documentos_referenciados": [doc.documento.nome for doc in documentos_similares]
            }

        except Exception as e:
            logger.error(f"Erro ao gerar resposta: {e}")
            # Atualiza a mensagem marcando como erro
            mensagem.resposta = f"Erro: {str(e)}"
            mensagem.erro = True
            mensagem.save()
            
            return {
                "success": False,
                "resposta": "Desculpe, ocorreu um erro ao processar sua solicitação.",
                "conversa_id": conversa.id,
                "error": str(e)
            }

    except Exception as e:
        logger.error(f"Erro no processo de busca de contexto: {e}")
        return {
            "success": False,
            "resposta": "Desculpe, ocorreu um erro ao iniciar a conversa.",
            "error": str(e)
        }